#!/usr/bin/env python3
"""
Simulation of the common ways agents edit .docx files.

Each mutator reproduces a real, documented pattern rather than an invented
worst case. The pattern being modelled is named in each docstring.

These are kept for regression coverage. The headline results now come from
real agent runs - see AGENT_RUNS.md.
"""
import re, os, shutil, zipfile, io

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _read_parts(path):
    with zipfile.ZipFile(path) as z:
        return {n: z.read(n) for n in z.namelist()}


def _write_parts(parts, path):
    with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as z:
        for n, d in parts.items():
            z.writestr(n, d)


# ---------------------------------------------------------------- A
def m_roundtrip_pythondocx(src, dst):
    """A. Open and save via python-docx, no edits.

    Baseline: how much breaks with no edit at all. This is what any agent
    using python-docx as a transport layer does implicitly.
    """
    import docx
    d = docx.Document(src)
    d.save(dst)


# ---------------------------------------------------------------- B
def m_pythondocx_settext(src, dst):
    """B. Replace paragraph text via paragraph.text = ...

    The most common way to "edit a document" in agent scripts and in library
    examples. Assigning .text drops every run in the paragraph and creates
    one new run - taking any anchors between runs with it.
    """
    import docx
    d = docx.Document(src)
    for p in d.paragraphs:
        if 'Effective Date' in p.text:
            p.text = p.text.replace('date of last signature',
                                    'date on which both parties have executed this agreement')
        if 'per annum' in p.text:
            p.text = p.text.replace('44,500', '47,250')
    d.save(dst)


# ---------------------------------------------------------------- C
def m_llm_raw_xml_value(src, dst):
    """C1. Targeted value replacement in raw XML.

    The agent was handed document.xml and changed a number by string
    replacement. The least destructive of the raw-XML patterns.
    """
    parts = _read_parts(src)
    x = parts['word/document.xml'].decode('utf-8')
    x = x.replace('EUR 44,500', 'EUR 47,250').replace('EUR 32,500', 'EUR 35,250')
    parts['word/document.xml'] = x.encode('utf-8')
    _write_parts(parts, dst)


def m_llm_copy_clause(src, dst):
    """C2. "Add another clause like the previous one."

    The agent copies a paragraph XML block and changes its text. The w:id
    revision attributes get copied along with it. This exact defect class is
    documented in the Python-Redlines README as a cause of Word's
    "unreadable content" warning.
    """
    parts = _read_parts(src)
    x = parts['word/document.xml'].decode('utf-8')
    m = re.search(r'<w:p>(?:(?!</w:p>).)*?<w:ins w:id="101".*?</w:p>', x, re.S)
    if m:
        block = m.group(0)
        clone = block.replace('professional indemnity insurance',
                              'cyber liability insurance of not less than EUR 5,000,000')
        x = x.replace(block, block + clone, 1)
    parts['word/document.xml'] = x.encode('utf-8')
    _write_parts(parts, dst)


def m_llm_prettyprint(src, dst):
    """C3. The agent "tidied up" the XML: reformatted it and dropped
    xml:space attributes it judged redundant.

    Observed behaviour when XML is put into context and written back.
    """
    parts = _read_parts(src)
    x = parts['word/document.xml'].decode('utf-8')
    x = x.replace(' xml:space="preserve"', '')
    x = re.sub(r'>\s+<', '><', x)
    parts['word/document.xml'] = x.encode('utf-8')
    _write_parts(parts, dst)


def m_llm_drop_paragraph(src, dst):
    """C4. "Drop the paragraph about historic spend."

    The agent deletes the whole paragraph without noticing it contains a
    footnote anchor. A textbook referential-integrity loss.
    """
    parts = _read_parts(src)
    x = parts['word/document.xml'].decode('utf-8')
    x = re.sub(r'<w:p><w:pPr><w:pStyle w:val="ClauseBody"/></w:pPr>\s*'
               r'<w:r><w:t xml:space="preserve">Historic spend.*?</w:p>', '', x, count=1, flags=re.S)
    parts['word/document.xml'] = x.encode('utf-8')
    _write_parts(parts, dst)


def m_llm_rename_style(src, dst):
    """C5. The agent renamed a style in styles.xml without updating the
    references in the body.

    Happens when style edits and content edits are separate steps.
    """
    parts = _read_parts(src)
    s = parts['word/styles.xml'].decode('utf-8')
    s = s.replace('w:styleId="ClauseBody"', 'w:styleId="BodyClause"')
    parts['word/styles.xml'] = s.encode('utf-8')
    _write_parts(parts, dst)


# ---------------------------------------------------------------- D
def m_markdown_roundtrip(src, dst):
    """D. Round-trip through markdown.

    The agent extracts text to markdown, edits it, and builds a new docx.
    A pattern practitioners describe openly as a way to "route around OOXML"
    (OOXML runs roughly 12x markdown in tokens).
    """
    import docx
    from docx import Document
    src_doc = Document(src)
    md_lines = []
    for block in src_doc.element.body.iterchildren():
        tag = block.tag.split('}')[-1]
        if tag == 'p':
            txt = ''.join(t.text or '' for t in block.iter(f'{{{W}}}t'))
            if txt.strip():
                md_lines.append(txt)
        elif tag == 'tbl':
            for tr in block.findall(f'{{{W}}}tr'):
                cells = [''.join(t.text or '' for t in tc.iter(f'{{{W}}}t'))
                         for tc in tr.findall(f'{{{W}}}tc')]
                md_lines.append(' | '.join(cells))
    out = Document()
    for line in md_lines:
        out.add_paragraph(line)
    out.save(dst)


MUTATORS = [
    ('A_roundtrip',      'python-docx: open and save, no edit',         m_roundtrip_pythondocx),
    ('B_settext',        'python-docx: paragraph.text = ...',           m_pythondocx_settext),
    ('C1_value',         'LLM edits a value in raw XML',                m_llm_raw_xml_value),
    ('C2_copyclause',    'LLM clones a block for "one more clause"',     m_llm_copy_clause),
    ('C3_prettyprint',   'LLM reformatted the XML',                     m_llm_prettyprint),
    ('C4_droppara',      'LLM deleted a para holding a footnote anchor', m_llm_drop_paragraph),
    ('C5_renamestyle',   'LLM renamed a style, left refs dangling',     m_llm_rename_style),
    ('D_markdown',       'round-trip through markdown',                 m_markdown_roundtrip),
]
