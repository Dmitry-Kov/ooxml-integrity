#!/usr/bin/env python3
"""Build and evaluate the versioned DOCX beta evidence corpus.

The labels in this file are ground truth declared before the checker runs.  A
mutation changes one known package property and names the findings that should
follow from that change.  Evaluation then compares the exact actionable
finding multiset with that label; it never records the checker's output as the
expected answer.

Source production is intentionally separate from output mutation.  The source
set is opened and saved by the producer named in the manifest.  Outputs are
then deterministic package-level edits of those committed sources.
"""
from __future__ import annotations

import argparse
import collections
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from lxml import etree

from ooxml_integrity import ERROR, Finding, Severity, __version__, check, compare


ROOT = Path(__file__).resolve().parents[1]
EVIDENCE = ROOT / "evidence" / "docx-beta"
SOURCES = EVIDENCE / "sources"
OUTPUTS = EVIDENCE / "outputs"
MANIFEST = EVIDENCE / "manifest.json"
METRICS = EVIDENCE / "metrics.json"
RESULTS = EVIDENCE / "RESULTS.md"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DCTERMS_NS = "http://purl.org/dc/terms/"
W = "{" + W_NS + "}"
REL = "{" + REL_NS + "}"
CP = "{" + CP_NS + "}"
DCTERMS = "{" + DCTERMS_NS + "}"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
ZIP_EPOCH = (2026, 1, 1, 0, 0, 0)
ACTIONABLE = {Severity.ERROR, Severity.WARN}

CATEGORIES = (
    "contract",
    "report",
    "letter",
    "table",
    "multi-section",
    "review-heavy",
)
PRODUCER_KEYS = ("python-docx", "libreoffice", "word-mac")

CATEGORY_TITLES = {
    "contract": "Services Agreement",
    "report": "Quarterly Operations Report",
    "letter": "Formal Project Letter",
    "table": "Asset Register",
    "multi-section": "Multi-section Policy Manual",
    "review-heavy": "Review Memorandum",
}


@dataclass(frozen=True)
class SourceSpec:
    id: str
    category: str
    ordinal: int
    producer: str


@dataclass(frozen=True)
class Expected:
    code: str
    severity: str
    count: int = 1

    def as_dict(self) -> dict[str, object]:
        return {
            "code": self.code,
            "severity": self.severity,
            "count": self.count,
        }


@dataclass(frozen=True)
class Mutation:
    description: str
    label_basis: str
    expected: tuple[Expected, ...]
    apply: Callable[[dict[str, bytes], list[str]], None]


def _source_specs() -> list[SourceSpec]:
    specs: list[SourceSpec] = []
    global_index = 0
    for category in CATEGORIES:
        for ordinal in range(1, 6):
            producer = PRODUCER_KEYS[global_index % len(PRODUCER_KEYS)]
            specs.append(SourceSpec(
                id=f"{category}-{ordinal:02d}",
                category=category,
                ordinal=ordinal,
                producer=producer,
            ))
            global_index += 1
    return specs


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _add_common_content(document: Document, spec: SourceSpec) -> None:
    marker = f"EVIDENCE-MARKER-{spec.id.upper()}-DRAFT"
    title = f"{CATEGORY_TITLES[spec.category]} {spec.ordinal}"

    document.add_heading(title, level=0)
    document.add_paragraph(marker, style="Subtitle")
    document.add_heading("Purpose and scope", level=1)
    document.add_paragraph(
        "This synthetic document exercises stable document structure without "
        "using customer text. It contains enough prose for meaningful fidelity "
        "volume checks and enough variation to avoid treating cloned bytes as "
        "distinct evidence. " * 4
    )
    review = document.add_paragraph(
        f"Reviewer sentence {spec.id}: confirm the stated obligations and the "
        "figures in the evidence table before approval."
    )
    document.add_comment(
        review.runs[0],
        text=f"Synthetic review note for {spec.id}.",
        author="Evidence Reviewer",
        initials="ER",
    )

    document.add_heading("Evidence table", level=1)
    table = document.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headings = ("Item", "Owner", "Status")
    for index, value in enumerate(headings):
        table.cell(0, index).text = value
    for row in range(1, 4):
        table.cell(row, 0).text = f"{spec.category} item {row}"
        table.cell(row, 1).text = f"Team {spec.ordinal}"
        table.cell(row, 2).text = ("Open", "Reviewed", "Approved")[row - 1]

    document.add_heading("Implementation notes", level=1)
    for item in range(1, 4):
        document.add_paragraph(
            f"Evidence step {item} for {spec.id} must remain readable and "
            "internally consistent after an automated edit.",
            style="List Number",
        )


def _add_category_content(document: Document, spec: SourceSpec) -> None:
    if spec.category == "contract":
        document.add_heading("Commercial terms", level=1)
        document.add_paragraph(
            f"The supplier shall deliver milestone {spec.ordinal} by the agreed "
            "date. The synthetic annual fee is EUR 10,000 and carries no real "
            "commercial meaning."
        )
    elif spec.category == "report":
        document.add_heading("Executive summary", level=1)
        document.add_paragraph(
            f"Synthetic throughput improved by {spec.ordinal + 10}% while the "
            "illustrative exception count remained below five."
        )
    elif spec.category == "letter":
        document.add_paragraph("1 Example Street\nSample City\n1 January 2026")
        document.add_paragraph("Dear Project Team,")
        document.add_paragraph(
            "Please treat this generated letter as test data only. It contains "
            "no actual names, addresses, commitments, or correspondence."
        )
        document.add_paragraph("Yours faithfully,\nEvidence Team")
    elif spec.category == "table":
        document.add_heading("Detailed register", level=1)
        table = document.add_table(rows=6, cols=4)
        table.style = "Table Grid"
        for row in range(6):
            for column in range(4):
                table.cell(row, column).text = (
                    f"Synthetic {spec.ordinal}-{row + 1}-{column + 1}"
                )
    elif spec.category == "multi-section":
        for section_number in (2, 3):
            section = document.add_section(WD_SECTION.NEW_PAGE)
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            section.header.paragraphs[0].text = (
                f"{spec.id} — section {section_number} header"
            )
            section.footer.paragraphs[0].text = (
                f"{spec.id} — section {section_number} footer"
            )
            document.add_heading(f"Section {section_number}", level=1)
            document.add_paragraph(
                f"This is the synthetic body of section {section_number}. " * 8
            )
    elif spec.category == "review-heavy":
        for note in range(2, 4):
            paragraph = document.add_paragraph(
                f"Review point {note} for {spec.id} needs an independent answer."
            )
            document.add_comment(
                paragraph.runs[0],
                text=f"Synthetic reviewer note {note} for {spec.id}.",
                author=f"Evidence Reviewer {note}",
                initials=f"E{note}",
            )


def _build_seed(spec: SourceSpec, path: Path) -> None:
    document = Document()
    if "CommentReference" not in document.styles:
        style = document.styles.add_style(
            "CommentReference", WD_STYLE_TYPE.CHARACTER,
        )
        style.font.size = Pt(8)

    props = document.core_properties
    fixed = datetime(2026, 1, 1, tzinfo=timezone.utc)
    props.title = f"{CATEGORY_TITLES[spec.category]} {spec.ordinal}"
    props.subject = "Synthetic OOXML integrity evidence"
    props.author = "ooxml-integrity evidence builder"
    props.keywords = f"synthetic,evidence,{spec.category}"
    props.comments = "No customer or personal data."
    props.created = fixed
    props.modified = fixed

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.header.paragraphs[0].text = f"{spec.id} — default header"
    section.footer.paragraphs[0].text = f"{spec.id} — default footer"

    _add_common_content(document, spec)
    _add_category_content(document, spec)
    document.save(path)


def _producer_versions(soffice: str | None, include_word: bool
                       ) -> dict[str, str]:
    import docx

    versions = {"python-docx": docx.__version__}
    if soffice:
        result = subprocess.run(
            [soffice, "--version"], check=True, capture_output=True, text=True,
        )
        versions["libreoffice"] = result.stdout.strip()
    if include_word:
        result = subprocess.run(
            ["osascript", "-e", 'tell application "Microsoft Word" to get version'],
            check=True, capture_output=True, text=True,
        )
        versions["word-mac"] = result.stdout.strip()
    return versions


def _word_save(source: Path, destination: Path) -> None:
    def quoted(path: Path) -> str:
        return str(path.resolve()).replace("\\", "\\\\").replace('"', '\\"')

    script = "\n".join((
        'tell application "Microsoft Word"',
        f'open file name "{quoted(source)}" add to recent files false',
        f'save as active document file name "{quoted(destination)}" '
        'file format format document add to recent files false',
        'close active document saving no',
        'end tell',
    ))
    subprocess.run(["osascript", "-e", script], check=True)


def _sanitise_source_metadata(path: Path) -> None:
    """Remove local Office identity/timestamps without changing document data."""
    parts, order = _read_package(path)
    core = _xml(parts, "docProps/core.xml")
    modified_by = core.find(CP + "lastModifiedBy")
    if modified_by is not None:
        modified_by.text = "ooxml-integrity evidence builder"
    for tag in ("created", "modified"):
        node = core.find(DCTERMS + tag)
        if node is not None:
            node.text = "2026-01-01T00:00:00Z"
    _store_xml(parts, "docProps/core.xml", core)
    _write_package(parts, order, path)


def _produce_sources(soffice: str | None, include_word: bool
                     ) -> tuple[list[dict[str, object]], dict[str, str]]:
    if not soffice:
        raise RuntimeError(
            "LibreOffice is required to build the complete source tranche; "
            "pass --soffice PATH"
        )
    if not include_word:
        raise RuntimeError(
            "Word for Mac is required to build the complete local source "
            "tranche; pass --word after reviewing the synthetic inputs"
        )

    SOURCES.mkdir(parents=True, exist_ok=True)
    versions = _producer_versions(soffice, include_word)
    records: list[dict[str, object]] = []
    with tempfile.TemporaryDirectory(prefix="ooxml-evidence-") as raw:
        staging = Path(raw)
        for spec in _source_specs():
            seed = staging / f"{spec.id}.docx"
            destination = SOURCES / seed.name
            _build_seed(spec, seed)
            if spec.producer == "python-docx":
                shutil.copyfile(seed, destination)
            elif spec.producer == "libreoffice":
                subprocess.run(
                    [
                        soffice,
                        "--headless",
                        "--convert-to", "docx",
                        "--outdir", str(SOURCES),
                        str(seed),
                    ],
                    check=True,
                    capture_output=True,
                    text=True,
                )
            elif spec.producer == "word-mac":
                _word_save(seed, destination)
            else:  # pragma: no cover - SourceSpec is internal and closed
                raise AssertionError(spec.producer)

            _sanitise_source_metadata(destination)
            actionable = [
                finding for finding in check(destination)
                if finding.severity in ACTIONABLE
            ]
            if actionable:
                raise RuntimeError(
                    f"producer source {spec.id} is not a clean control: "
                    f"{[(f.code, f.message) for f in actionable]}"
                )
            records.append({
                "id": spec.id,
                "path": f"sources/{destination.name}",
                "sha256": _sha256(destination),
                "category": spec.category,
                "producer": {
                    "id": spec.producer,
                    "version": versions[spec.producer],
                    "operation": "create" if spec.producer == "python-docx"
                    else "open-and-save as DOCX",
                },
                "synthetic": True,
                "personal_data": False,
                "postprocessing": (
                    "core lastModifiedBy/created/modified values sanitised and "
                    "ZIP container metadata normalised; document content and "
                    "relationships unchanged"
                ),
                "expected_behavior": (
                    "opens without a repair prompt; no actionable structural "
                    "finding is expected before mutation"
                ),
            })
    return records, versions


def _read_package(path: Path) -> tuple[dict[str, bytes], list[str]]:
    with zipfile.ZipFile(path) as archive:
        order = archive.namelist()
        return ({name: archive.read(name) for name in order}, order)


def _write_package(parts: dict[str, bytes], order: list[str], path: Path) -> None:
    with zipfile.ZipFile(
        path, "w", zipfile.ZIP_DEFLATED, compresslevel=9,
    ) as archive:
        for name in order:
            if name not in parts:
                continue
            info = zipfile.ZipInfo(name, ZIP_EPOCH)
            info.create_system = 3
            info.external_attr = (0o40775 if name.endswith("/") else 0o100664) << 16
            info.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(info, parts[name])
        for name in sorted(set(parts) - set(order)):
            info = zipfile.ZipInfo(name, ZIP_EPOCH)
            info.create_system = 3
            info.external_attr = 0o100664 << 16
            info.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(info, parts[name])


def _xml(parts: dict[str, bytes], part: str) -> etree._Element:
    try:
        return etree.fromstring(parts[part])
    except KeyError as exc:
        raise RuntimeError(f"fixture has no {part}") from exc


def _store_xml(parts: dict[str, bytes], part: str, root: etree._Element) -> None:
    parts[part] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True,
    )


def _find_text(root: etree._Element, needle: str | None = None
               ) -> etree._Element:
    for node in root.iter(W + "t"):
        if node.text and (needle is None or needle in node.text):
            return node
    raise RuntimeError(f"no text node containing {needle!r}")


def _safe_text_edit(parts: dict[str, bytes], _: list[str]) -> None:
    root = _xml(parts, "word/document.xml")
    node = _find_text(root, "DRAFT")
    node.text = (node.text or "").replace("DRAFT", "FINAL", 1)
    _store_xml(parts, "word/document.xml", root)


def _comment_id(root: etree._Element) -> str:
    reference = next(root.iter(W + "commentReference"), None)
    if reference is None:
        raise RuntimeError("source has no comment reference")
    value = reference.get(W + "id")
    if value is None:
        raise RuntimeError("comment reference has no id")
    return value


def _remove_matching(root: etree._Element, tag: str, comment_id: str) -> None:
    for node in list(root.iter(W + tag)):
        if node.get(W + "id") == comment_id:
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)


def _orphan_comment(parts: dict[str, bytes], _: list[str]) -> None:
    root = _xml(parts, "word/document.xml")
    comment_id = _comment_id(root)
    for tag in ("commentRangeStart", "commentRangeEnd", "commentReference"):
        _remove_matching(root, tag, comment_id)
    _store_xml(parts, "word/document.xml", root)


def _missing_comment_definition(parts: dict[str, bytes], _: list[str]) -> None:
    document = _xml(parts, "word/document.xml")
    comment_id = _comment_id(document)
    comments = _xml(parts, "word/comments.xml")
    _remove_matching(comments, "comment", comment_id)
    _store_xml(parts, "word/comments.xml", comments)


def _remove_comment_boundary(
        parts: dict[str, bytes], _: list[str], boundary: str) -> None:
    document = _xml(parts, "word/document.xml")
    comment_id = _comment_id(document)
    _remove_matching(document, boundary, comment_id)
    _store_xml(parts, "word/document.xml", document)


def _missing_comment_end(parts: dict[str, bytes], order: list[str]) -> None:
    _remove_comment_boundary(parts, order, "commentRangeEnd")


def _missing_comment_start(parts: dict[str, bytes], order: list[str]) -> None:
    _remove_comment_boundary(parts, order, "commentRangeStart")


def _change_story(parts: dict[str, bytes], kind: str) -> None:
    candidates = sorted(
        name for name in parts
        if name.startswith(f"word/{kind}") and name.endswith(".xml")
    )
    if not candidates:
        raise RuntimeError(f"source has no {kind} part")
    for part in candidates:
        root = _xml(parts, part)
        try:
            node = _find_text(root)
        except RuntimeError:
            # Word and LibreOffice may retain an empty story part alongside the
            # referenced story.  An empty implementation detail is not the
            # labelled story whose visible text this mutation is meant to lose.
            continue
        node.text = (node.text or "") + " [changed by seeded mutation]"
        _store_xml(parts, part, root)
        return
    raise RuntimeError(f"source has no non-empty {kind} story")


def _header_changed(parts: dict[str, bytes], _: list[str]) -> None:
    _change_story(parts, "header")


def _footer_changed(parts: dict[str, bytes], _: list[str]) -> None:
    _change_story(parts, "footer")


def _text_loss(parts: dict[str, bytes], _: list[str]) -> None:
    root = _xml(parts, "word/document.xml")
    for node in root.iter(W + "t"):
        node.text = ""
    _store_xml(parts, "word/document.xml", root)


def _broken_style(parts: dict[str, bytes], _: list[str]) -> None:
    root = _xml(parts, "word/document.xml")
    style = next(root.iter(W + "pStyle"), None)
    if style is None:
        raise RuntimeError("source has no paragraph style reference")
    style.set(W + "val", "MissingEvidenceStyle")
    _store_xml(parts, "word/document.xml", root)


def _whitespace_loss(parts: dict[str, bytes], _: list[str]) -> None:
    root = _xml(parts, "word/document.xml")
    node = _find_text(root, "DRAFT")
    node.text = " " + (node.text or "") + " "
    node.attrib.pop(XML_SPACE, None)
    _store_xml(parts, "word/document.xml", root)


def _missing_table_grid(parts: dict[str, bytes], _: list[str]) -> None:
    root = _xml(parts, "word/document.xml")
    table = next(root.iter(W + "tbl"), None)
    if table is None:
        raise RuntimeError("source has no table")
    grid = table.find(W + "tblGrid")
    if grid is None:
        raise RuntimeError("source table has no grid")
    table.remove(grid)
    _store_xml(parts, "word/document.xml", root)


def _table_row_mismatch(parts: dict[str, bytes], _: list[str]) -> None:
    root = _xml(parts, "word/document.xml")
    table = next(root.iter(W + "tbl"), None)
    if table is None:
        raise RuntimeError("source has no table")
    row = table.find(W + "tr")
    cells = [] if row is None else row.findall(W + "tc")
    if row is None or len(cells) < 2:
        raise RuntimeError("source first table row has too few cells")
    properties = cells[0].find(W + "tcPr")
    if properties is None:
        properties = etree.Element(W + "tcPr")
        cells[0].insert(0, properties)
    span = properties.find(W + "gridSpan")
    if span is None:
        span = etree.SubElement(properties, W + "gridSpan")
    span.set(W + "val", "2")
    _store_xml(parts, "word/document.xml", root)


def _missing_header_part(parts: dict[str, bytes], order: list[str]) -> None:
    relationships = _xml(parts, "word/_rels/document.xml.rels")
    relationship = next((
        item for item in relationships.iter(REL + "Relationship")
        if (item.get("Type") or "").endswith("/header")
    ), None)
    if relationship is None:
        raise RuntimeError("source has no header relationship")
    target = relationship.get("Target") or ""
    target = target.lstrip("/")
    part = target if target.startswith("word/") else "word/" + target
    if part not in parts:
        raise RuntimeError(f"header relationship target {part} is absent")
    parts.pop(part)
    if part in order:
        order.remove(part)


MUTATIONS: dict[str, Mutation] = {
    "clean-copy": Mutation(
        "byte-identical clean control",
        "The output is copied byte for byte; any actionable result is a false positive.",
        (),
        lambda _parts, _order: None,
    ),
    "safe-text-edit": Mutation(
        "same-length ordinary text edit",
        "Only the marker word DRAFT becomes FINAL; no supported structure is lost.",
        (),
        _safe_text_edit,
    ),
    "orphan-comment": Mutation(
        "remove one complete comment anchor but retain its body",
        "The comments part still defines the now-invisible note and fidelity loses one anchor.",
        (Expected("CMT005", "error"), Expected("FID001", "error")),
        _orphan_comment,
    ),
    "missing-comment-definition": Mutation(
        "remove one referenced comment body",
        "The document retains the reference while the labelled reviewer text is removed.",
        (Expected("CMT004", "error"), Expected("FID004", "error")),
        _missing_comment_definition,
    ),
    "header-changed": Mutation(
        "change text in one referenced header story",
        "The source header story no longer exists under any relationship or part name.",
        (Expected("FID007", "error"),),
        _header_changed,
    ),
    "footer-changed": Mutation(
        "change text in one referenced footer story",
        "The source footer story no longer exists under any relationship or part name.",
        (Expected("FID007", "error"),),
        _footer_changed,
    ),
    "text-loss": Mutation(
        "remove all main-story text while retaining its structure",
        "The edited main story is well below the declared 95% volume threshold.",
        (Expected("FID003", "error"),),
        _text_loss,
    ),
    "broken-style": Mutation(
        "point one paragraph at an undefined style",
        "Exactly one pStyle reference names a style absent from styles.xml.",
        (Expected("STY001", "error"),),
        _broken_style,
    ),
    "edge-whitespace": Mutation(
        "add edge whitespace without xml:space=preserve",
        "Exactly one text run contains edge whitespace that an OOXML consumer may eat.",
        (Expected("TXT001", "warn"),),
        _whitespace_loss,
    ),
    "missing-table-grid": Mutation(
        "remove the first table's grid",
        "The first table has no w:tblGrid after the mutation.",
        (Expected("TBL001", "error"),),
        _missing_table_grid,
    ),
    "table-row-mismatch": Mutation(
        "increase one cell span without changing the table grid",
        "The first row now spans one more column than w:tblGrid declares, "
        "without removing text or style references.",
        (Expected("TBL002", "warn"),),
        _table_row_mismatch,
    ),
    "missing-header-part": Mutation(
        "remove a referenced header package part",
        "The relationship target is absent and the requested fidelity comparison cannot complete.",
        (Expected("REL002", "error"), Expected("FID000", "error")),
        _missing_header_part,
    ),
    "missing-comment-end": Mutation(
        "remove one comment range end",
        "One commentRangeStart has no matching commentRangeEnd.",
        (Expected("CMT001", "error"),),
        _missing_comment_end,
    ),
    "missing-comment-start": Mutation(
        "remove one comment range start",
        "One commentRangeEnd has no matching commentRangeStart.",
        (Expected("CMT002", "error"),),
        _missing_comment_start,
    ),
}

CATEGORY_MUTATIONS = {
    "contract": ("orphan-comment", "missing-comment-definition"),
    "report": ("header-changed", "text-loss"),
    "letter": ("broken-style", "edge-whitespace"),
    "table": ("missing-table-grid", "table-row-mismatch"),
    "multi-section": ("footer-changed", "missing-header-part"),
    "review-heavy": ("missing-comment-end", "missing-comment-start"),
}


def _current_rule_inventory() -> list[str]:
    pattern = re.compile(r'["\']([A-Z]{3}\d{3})["\']')
    codes: set[str] = set()
    for path in (ROOT / "src" / "ooxml_integrity").glob("*.py"):
        codes.update(pattern.findall(path.read_text(encoding="utf-8")))
    # This is a DOCX corpus. PPTX rules have their own renderer-backed deck and
    # would be misleadingly listed as missing labels in this report.
    return sorted(code for code in codes if not code.startswith("PPT"))


def _build_outputs(sources: list[dict[str, object]], *,
                   evidence_root: Path = EVIDENCE) -> list[dict[str, object]]:
    outputs = evidence_root / "outputs"
    outputs.mkdir(parents=True, exist_ok=True)
    pairs: list[dict[str, object]] = []
    for source_record in sources:
        source_id = str(source_record["id"])
        source_path = evidence_root / str(source_record["path"])
        mutation_ids = (
            "clean-copy",
            "safe-text-edit",
            *CATEGORY_MUTATIONS[str(source_record["category"])],
        )
        for mutation_id in mutation_ids:
            mutation = MUTATIONS[mutation_id]
            pair_id = f"{source_id}--{mutation_id}"
            output_path = outputs / f"{pair_id}.docx"
            if mutation_id == "clean-copy":
                shutil.copyfile(source_path, output_path)
            else:
                parts, order = _read_package(source_path)
                mutation.apply(parts, order)
                _write_package(parts, order, output_path)
            pairs.append({
                "id": pair_id,
                "source": str(source_record["path"]),
                "output": f"outputs/{output_path.name}",
                "output_sha256": _sha256(output_path),
                "mutation": mutation_id,
                "class": "clean" if not mutation.expected else "seeded-defect",
                "description": mutation.description,
                "label_method": "isolated deterministic mutation",
                "label_basis": mutation.label_basis,
                "expected_findings": [item.as_dict() for item in mutation.expected],
            })
    return pairs


def build(soffice: str | None, include_word: bool) -> dict[str, object]:
    if MANIFEST.exists():
        raise RuntimeError("Refusing to replace a committed corpus. Use append-only producer import in this checkout.")
    if include_word and sys.platform != "darwin":
        raise RuntimeError("build --word is Mac-only; use prepare-windows and import-windows.")
    sources, versions = _produce_sources(soffice, include_word)
    pairs = _build_outputs(sources)
    manifest: dict[str, object] = {
        "schema_version": 1,
        "corpus": "docx-beta-synthetic-v1",
        "labelled_on": "2026-09-05",
        "scope": "DOCX structural integrity and source fidelity",
        "licence": "MIT (same as this repository)",
        "sanitisation": (
            "All documents are synthetic and contain no customer, personal, "
            "confidential, or externally licensed content."
        ),
        "label_policy": (
            "Expected findings are declared by isolated mutation before the "
            "checker runs. Exact actionable multisets are regression-gated."
        ),
        "known_evidence_gaps": [
            "No Word for Windows source has been collected.",
            "No Word Online source has been collected.",
            "No customer or independently supplied generator source is included.",
            "Labels have one maintainer review, not independent dual review.",
            "Renderer appearance is not scored by this structural corpus.",
        ],
        "producer_versions": versions,
        "rule_inventory": _current_rule_inventory(),
        "source_count": len(sources),
        "pair_count": len(pairs),
        "sources": sources,
        "pairs": pairs,
    }
    MANIFEST.write_text(
        json.dumps(manifest, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    return manifest


def rebuild_outputs(manifest_path: Path = MANIFEST) -> dict[str, object]:
    """Regenerate labelled edits without reopening committed producer sources."""
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    rebuilt = _build_outputs(manifest["sources"], evidence_root=manifest_path.parent)
    by_id = {pair["id"]: pair for pair in rebuilt}
    # An actual Office save cannot be reproduced by a package mutation. Keep
    # both its bytes and its receipt/label in their original manifest position.
    pairs = [by_id.pop(pair["id"]) if pair["mutation"] in MUTATIONS else pair
             for pair in manifest["pairs"]]
    pairs.extend(by_id.values())
    manifest["rule_inventory"] = _current_rule_inventory()
    manifest["pair_count"] = len(pairs)
    manifest["pairs"] = pairs
    manifest_path.write_text(
        json.dumps(manifest, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    return manifest


def _expected_counter(pair: dict[str, object]
                      ) -> collections.Counter[tuple[str, str]]:
    return collections.Counter({
        (str(item["code"]), str(item["severity"])): int(item["count"])
        for item in pair["expected_findings"]  # type: ignore[index]
    })


def _actual_findings(source: Path, output: Path) -> list[Finding]:
    findings = list(check(output))
    try:
        findings.extend(compare(source, output))
    except Exception as exc:
        findings.append(Finding(
            "FID000",
            ERROR,
            f"source comparison was requested but could not run: {exc}; "
            "comparison was NOT performed",
        ))
    return [finding for finding in findings if finding.severity in ACTIONABLE]


def _ratio(numerator: int, denominator: int) -> float | None:
    return numerator / denominator if denominator else None


def evaluate(manifest_path: Path = MANIFEST, *, verify_hashes: bool = True
             ) -> dict[str, object]:
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    root = manifest_path.parent
    failures: list[dict[str, object]] = []
    expected_by_rule: collections.Counter[str] = collections.Counter()
    actual_by_rule: collections.Counter[str] = collections.Counter()
    true_by_rule: collections.Counter[str] = collections.Counter()
    expected_errors = actual_errors = true_errors = 0

    source_hashes = {
        str(item["path"]): str(item["sha256"])
        for item in manifest["sources"]
    }
    artifact_hashes = {str(item["path"]): str(item["sha256"])
                       for item in manifest.get("supporting_artifacts", [])}
    producer_by_path = {str(item["path"]): item["producer"]["id"] for item in manifest["sources"]}
    groups: dict[str, dict[str, object]] = {}
    if verify_hashes:
        for relative, expected_hash in {**source_hashes, **artifact_hashes}.items():
            path = root / relative
            actual_hash = _sha256(path)
            if actual_hash != expected_hash:
                raise RuntimeError(
                    f"source hash mismatch for {relative}: "
                    f"{actual_hash} != {expected_hash}"
                )

    for pair in manifest["pairs"]:
        source = root / pair["source"]
        output = root / pair["output"]
        registered_hash = source_hashes.get(pair["source"], artifact_hashes.get(pair["source"]))
        if registered_hash is None:
            raise RuntimeError(f"unregistered pair source: {pair['source']}")
        if pair.get("source_sha256", registered_hash) != registered_hash:
            raise RuntimeError(f"pair source hash disagrees with artifact: {pair['id']}")
        if verify_hashes:
            actual_hash = _sha256(output)
            if actual_hash != pair["output_sha256"]:
                raise RuntimeError(
                    f"output hash mismatch for {pair['output']}: "
                    f"{actual_hash} != {pair['output_sha256']}"
                )
        expected = _expected_counter(pair)
        actual_findings = _actual_findings(source, output)
        actual = collections.Counter(
            (finding.code, finding.severity.value)
            for finding in actual_findings
        )
        producer = producer_by_path.get(pair["source"], producer_by_path.get(pair["output"]))
        if producer is None:
            raise RuntimeError(f"unregistered producer for pair: {pair['id']}")
        kind = "word-roundtrip" if pair["mutation"] == "word-roundtrip" else "deterministic-mutation"
        for group_key in (f"producer:{producer}", f"kind:{kind}"):
            group = groups.setdefault(group_key, {"pairs": 0, "clean_pairs": 0, "tp": 0, "fp": 0, "fn": 0})
            group["pairs"] += 1
            group["clean_pairs"] += pair["class"] == "clean"
            for key in set(expected) | set(actual):
                if key[1] == "error":
                    matched = min(expected[key], actual[key])
                    group["tp"] += matched
                    group["fp"] += actual[key] - matched
                    group["fn"] += expected[key] - matched
        if actual != expected:
            failures.append({
                "pair": pair["id"],
                "expected": [
                    {"code": key[0], "severity": key[1], "count": count}
                    for key, count in sorted(expected.items())
                ],
                "actual": [
                    {"code": key[0], "severity": key[1], "count": count}
                    for key, count in sorted(actual.items())
                ],
            })

        for (code, severity), count in expected.items():
            expected_by_rule[code] += count
            if severity == "error":
                expected_errors += count
        for (code, severity), count in actual.items():
            actual_by_rule[code] += count
            if severity == "error":
                actual_errors += count
        for key in set(expected) | set(actual):
            matched = min(expected[key], actual[key])
            true_by_rule[key[0]] += matched
            if key[1] == "error":
                true_errors += matched

    inventory = sorted(set(manifest["rule_inventory"]) | set(expected_by_rule)
                       | set(actual_by_rule))
    rules = []
    for code in inventory:
        expected = expected_by_rule[code]
        actual = actual_by_rule[code]
        matched = true_by_rule[code]
        false_positive = actual - matched
        false_negative = expected - matched
        rules.append({
            "code": code,
            "status": "measured" if expected or actual else "not-measured",
            "tp": matched,
            "fp": false_positive,
            "fn": false_negative,
            "precision": _ratio(matched, matched + false_positive),
            "recall": _ratio(matched, matched + false_negative),
        })

    error_fp = actual_errors - true_errors
    error_fn = expected_errors - true_errors
    for group in groups.values():
        group["precision"] = _ratio(group["tp"], group["tp"] + group["fp"])
        group["recall"] = _ratio(group["tp"], group["tp"] + group["fn"])
    return {
        "schema_version": 1,
        "corpus": manifest["corpus"],
        "checker_version": __version__,
        "sources": len(manifest["sources"]),
        "pairs": len(manifest["pairs"]),
        "clean_pairs": sum(
            pair["class"] == "clean" for pair in manifest["pairs"]
        ),
        "seeded_defect_pairs": sum(
            pair["class"] == "seeded-defect" for pair in manifest["pairs"]
        ),
        "error_level": {
            "tp": true_errors,
            "fp": error_fp,
            "fn": error_fn,
            "precision": _ratio(true_errors, true_errors + error_fp),
            "recall": _ratio(true_errors, true_errors + error_fn),
        },
        "rules": rules,
        "groups": groups,
        "pair_failures": failures,
    }


def _percent(value: float | None) -> str:
    return "not measured" if value is None else f"{value * 100:.1f}%"


def _write_results(metrics: dict[str, object]) -> None:
    METRICS.write_text(
        json.dumps(metrics, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    error = metrics["error_level"]
    measured = [item for item in metrics["rules"] if item["status"] == "measured"]
    unmeasured = [
        item["code"] for item in metrics["rules"]
        if item["status"] == "not-measured"
    ]
    lines = [
        "# DOCX beta evidence results",
        "",
        "This report is generated from `manifest.json` by "
        "`research/build_docx_evidence.py evaluate --write`. The manifest's "
        "expected labels are declared by isolated mutations; they are not "
        "snapshots of checker output.",
        "",
        "## Corpus denominator",
        "",
        f"- Sources: **{metrics['sources']}**.",
        f"- Labelled source/output pairs: **{metrics['pairs']}**.",
        f"- Clean controls: **{metrics['clean_pairs']}**.",
        f"- Seeded-defect pairs: **{metrics['seeded_defect_pairs']}**.",
        "- Unit of counting: one actionable finding occurrence. Exact duplicate "
        "counts matter; info-level observations are outside this precision gate.",
        "",
        "## Error-level result",
        "",
        f"- True positives: **{error['tp']}**.",
        f"- False positives: **{error['fp']}**.",
        f"- False negatives: **{error['fn']}**.",
        f"- Precision: **{_percent(error['precision'])}** "
        "(`TP / (TP + FP)`).",
        f"- Recall: **{_percent(error['recall'])}** "
        "(`TP / (TP + FN)`).",
        "",
        "## Rule-level result",
        "",
        "| rule | TP | FP | FN | precision | recall |",
        "| --- | ---: | ---: | ---: | ---: | ---: |",
    ]
    for item in measured:
        lines.append(
            f"| `{item['code']}` | {item['tp']} | {item['fp']} | "
            f"{item['fn']} | {_percent(item['precision'])} | "
            f"{_percent(item['recall'])} |"
        )
    lines.extend(("", "## Producer and pair-kind error results", "",
                  "| group | pairs | clean | TP | FP | FN | precision | recall |",
                  "| --- | ---: | ---: | ---: | ---: | ---: | ---: | ---: |"))
    for name, group in sorted(metrics["groups"].items()):
        lines.append(f"| `{name}` | {group['pairs']} | {group['clean_pairs']} | {group['tp']} | {group['fp']} | {group['fn']} | {_percent(group['precision'])} | {_percent(group['recall'])} |")
    lines.extend((
        "",
        "Rules with no positive or negative label in this tranche are explicitly "
        "not measured; silence is not treated as evidence:",
        "",
        ", ".join(f"`{code}`" for code in unmeasured) + ".",
        "",
        "## Interpretation boundary",
        "",
        "These numbers establish reproducible regression behaviour on synthetic "
        "DOCX package mutations. They do **not** establish production precision "
        "for unmeasured rules, customer document distributions, other Windows "
        "Word versions, Word Online, or visual renderer fidelity. Those gaps are kept "
        "in `manifest.json` and the corpus README rather than being folded into "
        "the 100% measured-rule result.",
        "",
    ))
    RESULTS.write_text("\n".join(lines), encoding="utf-8")


def _check_floor(metrics: dict[str, object]) -> list[str]:
    problems = []
    if metrics["sources"] < 30:
        problems.append("fewer than 30 sources")
    if metrics["pairs"] < 100:
        problems.append("fewer than 100 labelled pairs")
    precision = metrics["error_level"]["precision"]
    if precision is None or precision < 0.95:
        problems.append("error-level precision is below 95%")
    if metrics["pair_failures"]:
        problems.append(f"{len(metrics['pair_failures'])} pair label mismatch(es)")
    return problems


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    commands = parser.add_subparsers(dest="command", required=True)
    build_parser = commands.add_parser("build", help="build sources, outputs and labels")
    build_parser.add_argument(
        "--soffice", default=shutil.which("soffice"),
        help="LibreOffice executable used for the producer tranche",
    )
    build_parser.add_argument(
        "--word", action="store_true",
        help="allow Microsoft Word for Mac to open and save its synthetic tranche",
    )
    commands.add_parser(
        "rebuild-outputs",
        help="rebuild mutations and labels from committed producer sources",
    )
    for name in ("prepare-windows", "import-windows"):
        windows_parser = commands.add_parser(name, help="append-only Windows Word evidence workflow")
        windows_parser.add_argument("--staging", type=Path, required=True)
    evaluate_parser = commands.add_parser("evaluate", help="score committed labels")
    evaluate_parser.add_argument(
        "--write", action="store_true", help="update metrics.json and RESULTS.md",
    )
    args = parser.parse_args(argv)

    if args.command in {"prepare-windows", "import-windows"}:
        try:
            from . import windows_docx_evidence as windows
        except ImportError:
            import windows_docx_evidence as windows
        if args.command == "prepare-windows":
            request = windows.prepare(args.staging)
            print(f"prepared {len(request['documents'])} clean inputs; next run save_docx_word_windows.ps1 -Batch {args.staging / 'batch.json'} in the desktop user session")
            return 0
        manifest = windows.import_batch(args.staging)
        print(f"imported Windows evidence: {manifest['source_count']} sources, {manifest['pair_count']} pairs")
        metrics = evaluate()
        _write_results(metrics)
    elif args.command == "build":
        manifest = build(args.soffice, args.word)
        print(
            f"built {manifest['source_count']} sources and "
            f"{manifest['pair_count']} labelled pairs"
        )
        metrics = evaluate()
        _write_results(metrics)
    elif args.command == "rebuild-outputs":
        manifest = rebuild_outputs()
        print(f"rebuilt {manifest['pair_count']} labelled outputs")
        metrics = evaluate()
        _write_results(metrics)
    else:
        metrics = evaluate()
        if args.write:
            _write_results(metrics)

    problems = _check_floor(metrics)
    error = metrics["error_level"]
    print(
        f"sources={metrics['sources']} pairs={metrics['pairs']} "
        f"error precision={_percent(error['precision'])} "
        f"recall={_percent(error['recall'])} "
        f"label mismatches={len(metrics['pair_failures'])}"
    )
    if problems:
        for problem in problems:
            print(f"evidence: {problem}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
